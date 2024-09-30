from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def gerar_relatorio_word(livros, caminho_word):
    # Criação do documento Word
    doc = Document()

    # Título do Relatório
    titulo = doc.add_heading("Relatório de Livros", level=1)
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Adiciona uma tabela para os dados dos livros
    tabela = doc.add_table(rows=1, cols=5)
    tabela.style = 'Table Grid'

    # Definindo os títulos da tabela
    hdr_cells = tabela.rows[0].cells
    hdr_cells[0].text = 'ID'
    hdr_cells[1].text = 'Título'
    hdr_cells[2].text = 'Autor'
    hdr_cells[3].text = 'Ano'
    hdr_cells[4].text = 'Preço'

    # Preenchendo a tabela com os dados dos livros
    for livro in livros:
        row_cells = tabela.add_row().cells
        row_cells[0].text = str(livro[0])
        row_cells[1].text = livro[1]
        row_cells[2].text = livro[2]
        row_cells[3].text = str(livro[3])
        row_cells[4].text = f"R$ {livro[4]:.2f}"

    # Salva o documento
    doc.save(caminho_word)
    print(f"Relatório Word gerado com sucesso em: {caminho_word}")
